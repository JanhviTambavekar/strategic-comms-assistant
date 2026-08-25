"""Extracts plain text from uploaded .txt, .pdf, .md, or .docx files."""
from io import BytesIO


def extract_text(filename: str, data: bytes) -> str:
    """Return extracted text from an uploaded file's raw bytes."""
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return _extract_pdf(data)
    if name.endswith(".docx"):
        return _extract_docx(data)
    # default: treat as text
    try:
        return data.decode("utf-8", errors="ignore").strip()
    except Exception:
        return ""


def _extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        return "[pypdf not installed - cannot read PDF]"
    try:
        reader = PdfReader(BytesIO(data))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages).strip()
    except Exception as exc:
        return f"[Could not extract PDF text: {exc}]"


def _extract_docx(data: bytes) -> str:
    """Extract text from a .docx file using python-docx."""
    try:
        import docx
    except ImportError:
        return "[python-docx not installed - cannot read .docx files]"
    try:
        document = docx.Document(BytesIO(data))
        parts = []
        for para in document.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                parts.append(" | ".join(cells))
        return "\n".join(parts).strip()
    except Exception as exc:
        return f"[Could not extract DOCX text: {exc}]"
