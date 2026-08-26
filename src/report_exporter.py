"""Create downloadable Word and PDF versions of a Markdown strategy report."""
from io import BytesIO
import re
from xml.sax.saxutils import escape


def _plain(text: str) -> str:
    text = re.sub(r"\[([^]]+)]\([^)]+\)", r"\1", text)
    return re.sub(r"[*_`]+", "", text).strip()


def _blocks(markdown_text: str):
    lines = markdown_text.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].strip()
        if not line:
            index += 1
            continue
        if line.startswith("|") and line.endswith("|"):
            rows = []
            while index < len(lines):
                candidate = lines[index].strip()
                if not (candidate.startswith("|") and candidate.endswith("|")):
                    break
                cells = [_plain(cell) for cell in candidate.strip("|").split("|")]
                separator = all(
                    re.fullmatch(r":?-{3,}:?", cell.replace(" ", ""))
                    for cell in cells
                )
                if not separator:
                    rows.append(cells)
                index += 1
            if rows:
                yield "table", rows
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            yield "heading", (len(heading.group(1)), _plain(heading.group(2)))
        elif re.match(r"^[-*]\s+", line):
            yield "bullet", _plain(re.sub(r"^[-*]\s+", "", line))
        elif re.match(r"^\d+[.)]\s+", line):
            yield "number", _plain(re.sub(r"^\d+[.)]\s+", "", line))
        else:
            yield "paragraph", _plain(line)
        index += 1


def to_docx(markdown_text: str) -> bytes:
    from docx import Document
    from docx.shared import Inches, Pt

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    document.styles["Normal"].font.name = "Aptos"
    document.styles["Normal"].font.size = Pt(10.5)
    for kind, value in _blocks(markdown_text):
        if kind == "heading":
            level, text = value
            document.add_heading(text, level=min(level, 3))
        elif kind == "bullet":
            document.add_paragraph(value, style="List Bullet")
        elif kind == "number":
            document.add_paragraph(value, style="List Number")
        elif kind == "table":
            width = max(len(row) for row in value)
            table = document.add_table(rows=0, cols=width)
            table.style = "Table Grid"
            for row_index, row in enumerate(value):
                cells = table.add_row().cells
                for column, text in enumerate(row):
                    cells[column].text = text
                    if row_index == 0:
                        for run in cells[column].paragraphs[0].runs:
                            run.bold = True
        else:
            document.add_paragraph(value)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def to_pdf(markdown_text: str) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    output = BytesIO()
    document = SimpleDocTemplate(
        output, pagesize=A4, rightMargin=16 * mm, leftMargin=16 * mm,
        topMargin=15 * mm, bottomMargin=15 * mm,
        title="Strategic Communications Strategy",
    )
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name="ReportTitle", parent=styles["Title"], fontSize=18, leading=22,
        textColor=colors.HexColor("#16324F"), spaceAfter=12,
    ))
    story = []
    for kind, value in _blocks(markdown_text):
        if kind == "heading":
            level, text = value
            style = "ReportTitle" if level == 1 else "Heading2" if level == 2 else "Heading3"
            story.extend([Paragraph(escape(text), styles[style]), Spacer(1, 4)])
        elif kind in {"bullet", "number"}:
            prefix = "• " if kind == "bullet" else ""
            story.extend([Paragraph(escape(prefix + value), styles["BodyText"]), Spacer(1, 3)])
        elif kind == "table":
            rows = [[Paragraph(escape(cell), styles["BodyText"]) for cell in row] for row in value]
            table = Table(rows, repeatRows=1, hAlign="LEFT")
            table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#DCEAF3")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#8FA8B8")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]))
            story.extend([table, Spacer(1, 8)])
        else:
            story.extend([Paragraph(escape(value), styles["BodyText"]), Spacer(1, 5)])
    document.build(story)
    return output.getvalue()
